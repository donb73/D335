from docx import Document
from docx.shared import Pt

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

doc.add_heading('Resume', 0)
doc.add_paragraph('[Your Full Name]\n[City, State] • [Email] • [Phone] • [LinkedIn]')

doc.add_heading('Professional Summary', level=1)
doc.add_paragraph("Detail-oriented and self-directed QA professional with over a decade of experience in software quality assurance, manual testing, and test documentation across highly regulated and fast-paced environments...")

doc.add_heading('Technical Skills', level=1)
doc.add_paragraph("Testing Tools: Jira, Zephyr, HP ALM, Bugzilla, Postman\n"
                  "Documentation & Platforms: Confluence, Microsoft Excel (advanced), Git, Octopus Deploy\n"
                  "Databases & Languages: SQL, Microsoft SQL Server Management Studio, C# (basic), Python (basic)\n"
                  "Monitoring & Reporting: Application Insights, Datadog, Grafana\n"
                  "Methods: Agile/Scrum, Risk-Based Testing, Exploratory Testing")

doc.add_heading('Professional Experience', level=1)

experience = [
    ("R1 RCM – Boise, ID", "Quality Assurance Analyst II | 2021–2025", [
        "Independently managed test queues...",
        "Maintained thorough test evidence...",
    ]),
    ("VisitPay – Boise, ID", "Quality Assurance Engineer | 2018–2021", [
        "Led test strategy...",
        "Applied risk-based testing...",
    ]),
    ("HealthCast – Boise, ID", "Software Quality Tester | 2016–2018", [
        "Standardized bug tracking workflows using Bugzilla...",
    ]),
    ("Beyondsoft – Boise, ID", "Test Engineer III | 2012–2016", [
        "Created detailed test plans...",
    ]),
    ("QualityLogic – Boise, ID", "QA Test Lead | 2002–2012", [
        "Led QA teams for multiple test sessions...",
        "Designed advanced Excel-based test systems...",
    ])
]

for company, title, bullets in experience:
    doc.add_heading(company, level=2)
    doc.add_paragraph(title, style='Intense Quote')
    for bullet in bullets:
        doc.add_paragraph(f'• {bullet}', style='List Bullet')

doc.add_heading('Certifications', level=1)
doc.add_paragraph("• Microsoft Certified: Azure AI Fundamentals – 2024\n"
                  "• Microsoft Certified: Azure Fundamentals – 2024\n"
                  "• Microsoft Certified: Azure Data Fundamentals – 2024")

doc.add_heading('Additional Qualifications', level=1)
doc.add_paragraph("• U.S. Citizen – eligible for roles requiring federal security clearance\n"
                  "• Willing to travel up to 20% as needed\n"
                  "• Flexible and adaptable to new internal tools and testing platforms")

doc.save("Oracle_Configuration_Test_Analyst_Resume.docx")
